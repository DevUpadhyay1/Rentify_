from docx import Document
from docx.shared import Pt

report_sections = [
    ("Project Title", "Rentify"),
    ("Group Members (fill names below)", "1. Member 1: __________________ - Backend Lead\n2. Member 2: __________________ - Frontend Lead\n3. Member 3: __________________ - QA / PM\n4. Member 4: __________________ - DevOps & Documentation"),
    ("Chapter 1: Introduction", \
     "1.1 Project Summary\nRentify is a web app where people can list items for rent and others can rent them. Owners add item details and pictures. Renters search, book, pay, and review. The backend uses Django and the frontend uses React with Vite.\n\n1.2 Purpose\nTo let people lend and borrow items easily, making it simple to earn and save money without buying things.\n\n1.3 Objective\n- Let owners add and manage items.\n- Let renters search, book, and pay.\n- Keep bookings from clashing.\n- Support user accounts and reviews.\n- Make the app safe and easy to use.\n\n1.4 Scope\nIncluded:\n- User signup/login, profile.\n- Item listing (with images), categories.\n- Search and basic filters.\n- Booking flow and booking status tracking.\n- Reviews and ratings.\n- Admin abilities through Django admin.\n\nNot included (for now):\n- Production-grade payment flow (Razorpay keys are present in settings, but full flow may need wiring).\n- Mobile app (but site is responsive).\n- Full deployment automation (Docker/gunicorn/nginx not present).\n\n1.5 Technology and short literature notes\n- Backend: Django 5.2, Django REST Framework, Simple JWT.\n- Frontend: React + Vite, Tailwind, Axios.\n- Database: Postgres configured in settings (local dev can use SQLite).\n- Storage: Firebase is used for image uploads.\n- Payments: Razorpay credentials expected via env vars."),
    ("Chapter 2: Project Management", \
     "2.1 Project Planning\nPhases: planning (2 weeks), development (6 weeks), testing and deployment (2 weeks).\n\n2.1.1 Development approach and justification\nIncremental development (Agile-like). Build core features first, then add extras.\n\n2.1.2 Effort and cost (simple estimate)\n- Team: 4 people (Backend, Frontend, QA/PM, DevOps/Docs).\n- Time: ~10 weeks.\n\n2.1.3 Roles and responsibilities\n- Backend: APIs, models, authentication.\n- Frontend: UI, routes, forms.\n- QA/PM: testing, user acceptance, coordinate tasks.\n- DevOps/Docs: deployment, environment setup, documentation.\n\n2.1.4 Group dependencies\n- Frontend depends on backend API stability.\n- Payments depend on selecting and configuring a payment gateway.\n- Notifications depend on Firebase/email.\n\n2.2 Project scheduling\n- Weeks 1–2: auth, basic item model, frontend skeleton.\n- Weeks 3–6: listing, searching, booking flows.\n- Weeks 7–8: payments, reviews, admin polish.\n- Weeks 9–10: testing, bug fixes, deploy prep."),
    ("Chapter 3: System Requirements Study", \
     "3.1 User Characteristics\n- Owners: create item listings, set price and availability.\n- Renters: search, book, pay, review.\n- Admin: manage data via Django admin.\n\n3.2 Hardware and Software Requirements\nDevelopment:\n- Python 3.10+ and virtualenv.\n- Node.js (for frontend), Vite.\n- PostgreSQL for production; SQLite possible for quick dev.\n- Modern browser for UI.\n\nProduction (small):\n- Linux server with 1–2 vCPUs, 2–4 GB RAM, Postgres, object storage or Firebase, HTTPS.\n\n3.3 Assumptions and Dependencies\n- Environment variables for email, Firebase, Razorpay must be set.\n- Users have internet access.\n- Third-party services (Firebase, Razorpay, email provider) are available."),
    ("Chapter 4: System Analysis", \
     "4.1 Study of Current System\n- Auth: custom user model and email verification.\n- Items: categories, items, images, wishlist.\n- Rental: booking and booking history with price calc and statuses.\n- Review: rating code exists and is referenced by models.\n\n4.2 Problems and Weaknesses\n- Payment flow may not be fully implemented.\n- Missing env vars will break email/Firebase/payments if not set.\n- Frontend serving via Django commented in settings.\n- Concurrency: availability checks exist but need DB-level locks for scale.\n\n4.3 Requirements of New System\n- Functional: register/verify, login, item CRUD, image upload, search, booking, payment handling, reviews.\n- Non-functional: security (JWT), performance, maintainability.\n\n4.4 Activity / Process\n- Owner creates item.\n- Renter requests booking.\n- System checks availability.\n- Booking created pending -> payment -> confirmed.\n\n4.5 Integration\n- Firebase, email SMTP, Razorpay.\n\n4.6 Features\n- Listings, calendar availability, booking, payments, reviews, admin.\n\n4.7 Use Case and Models\n- Users, Items, Bookings, Reviews, Payments (expected)."),
    ("Chapter 5: System Design", \
     "5.1 Application Design\n- Input: HTTP API requests and files.\n- Output: JSON responses and emails.\n\n5.1.1 Pseudocode (booking)\n- Check user auth.\n- Check item availability.\n- Calculate price in Booking.save().\n- Create booking pending.\n- Trigger payment flow.\n- On success: confirm booking.\n\n5.2 API and Interfaces\n- Auth endpoints (register, verify, login, profile).\n- Item endpoints with image upload.\n- Booking endpoints.\n\n5.3 Security\n- JWT auth, token blacklist, image validation, use HTTPS in production."),
    ("Chapter 6: Implementation Planning", \
     "6.1 Environment\n- Activate virtualenv in project.\n- Install Python and Node dependencies.\n- Set env vars for email/Firebase/payments.\n\n6.2 Module spec\n- Authentication app: register/login/verify/profile.\n- Items app: categories, items, images, wishlist.\n- Rental app: booking and history.\n- billing app: expected for payments.\n\n6.3 Security Features\n- Password hashing, JWT, CSRF for forms, input validation, file size/type checks.\n\n6.4 Coding Standard\n- Python: PEP8.\n- JS: ESLint.\n"),
    ("Chapter 7: Testing", \
     "7.1 Testing Plan\n- Unit tests for models and business logic.\n- API tests for endpoints.\n- Frontend component tests.\n\n7.2 Strategy\n- Prioritize auth and booking flows.\n- Use CI for tests.\n\n7.3 Test suites examples\n- Auth, Items, Booking, Review tests."),
    ("Chapter 8: Conclusion and Discussion", \
     "8.1 Viability\n- Core features are implemented and project is viable.\n\n8.2 Problems & Solutions\n- Double booking: add DB transactions and locks.\n- Missing env vars: provide .env.example and docs.\n- Image storage: consider S3 for production.\n\n8.3 Summary\n- Backend is mostly complete; frontend skeleton exists. Next: payments, deployment, hardening."),
    ("Chapter 9: Limitations and Future Enhancements", \
     "Limitations:\n- Payment flow may be incomplete.\n- No Docker or CI files present.\n\nFuture enhancements:\n- Implement payment gateway fully.\n- Add Docker and deploy pipeline.\n- Improve booking transaction handling and analytics."),
    ("How to run locally", \
     "1. Activate virtualenv on Windows PowerShell: .\\env\\Scripts\\Activate.ps1\n2. Install Python requirements: pip install -r requirements.txt (if present)\n3. Frontend: cd r_frontend; npm install; npm run dev\4. Django: python manage.py migrate; python manage.py runserver\n5. Set env vars for email, Firebase, Razorpay before starting.")
]


def add_paragraph(doc, text, bold=False):
    # sanitize text to remove control characters not allowed in XML
    if text is None:
        text = ''
    clean = ''.join(ch for ch in text if (ch == '\n' or ch == '\t' or ord(ch) >= 32))
    p = doc.add_paragraph()
    run = p.add_run(clean)
    run.bold = bold
    run.font.size = Pt(11)


def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Title
    doc.add_heading('Rentify - Project Report', level=0)

    for title, body in report_sections:
        doc.add_heading(title, level=1)
        # Split into paragraphs
        for para in body.split('\n\n'):
            # sanitize each paragraph before adding
            if para is None:
                para = ''
            add_paragraph(doc, para)

    output_path = 'c:/Users/Dev/Desktop/Rentify 2/Rentify_Report.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    main()
